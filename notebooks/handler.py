# OCR 파트
import json
import time
import requests
from pathlib import Path
import re
import fitz  # PyMuPDF

import os
from dotenv import load_dotenv

# 파일별 파싱 함수
from pathlib import Path
import pandas as pd
from docx import Document
import pdfplumber

load_dotenv()

INVOKE_URL = os.getenv("CLOVA_INVOKE_URL")
SECRET_KEY = os.getenv("CLOVA_OCR_SECRET")

# 파일(이미지/페이지)을 bytes 형태로 Clova OCR API에 보내서 OCR 결과 JSON을 받아옴
def _call_clova_from_bytes(
    filename: str,
    file_bytes: bytes,
    fmt: str,
    lang: str = "ko",
    enable_table: bool = True,
    timeout: int = 180,
) -> dict:
    headers = {"X-OCR-SECRET": SECRET_KEY}
    message = {
        "version": "V1",
        "requestId": f"clova-{Path(filename).stem}-{int(time.time()*1000)}",
        "timestamp": int(time.time() * 1000),
        "lang": lang,
        "images": [{"format": fmt, "name": Path(filename).stem}],
        "enableTableDetection": bool(enable_table),
    }
    files = {
        "file": (filename, file_bytes, "application/octet-stream"),
        "message": (None, json.dumps(message), "application/json"),
    }
    r = requests.post(INVOKE_URL, headers=headers, files=files, timeout=timeout)
    r.raise_for_status()
    return r.json()

# PDF를 페이지별로 렌더링해서 PNG 이미지 바이트로 변환
def _pdf_to_png_bytes_list(pdf_path: str, dpi: int = 200) -> list[tuple[str, bytes]]:
    p = Path(pdf_path)
    doc = fitz.open(str(p))
    zoom = dpi / 72
    mat = fitz.Matrix(zoom, zoom)

    out: list[tuple[str, bytes]] = []
    for i in range(len(doc)):
        pix = doc[i].get_pixmap(matrix=mat, alpha=False)
        out.append((f"{p.stem}_p{i+1}.png", pix.tobytes("png")))
    doc.close()
    return out

# Clova가 준 테이블 셀 1개에서 텍스트를 뽑는데, 셀 안에 여러 줄이면 줄바꿈(\n)을 유지해서 합침
def _cell_text_keep_lines(cell: dict) -> str:
    lines: list[str] = []

    if isinstance(cell.get("cellTextLines"), list) and cell["cellTextLines"]:
        for ln in cell["cellTextLines"]:
            if isinstance(ln.get("cellWords"), list) and ln["cellWords"]:
                s = " ".join(
                    [w.get("inferText", "") for w in ln["cellWords"] if w.get("inferText")]
                )
            else:
                s = ln.get("inferText", "")
            s = re.sub(r"\s+", " ", str(s)).strip()
            if s:
                lines.append(s)

    elif isinstance(cell.get("cellWords"), list) and cell["cellWords"]:
        s = " ".join([w.get("inferText", "") for w in cell["cellWords"] if w.get("inferText")])
        s = re.sub(r"\s+", " ", str(s)).strip()
        if s:
            lines.append(s)

    else:
        s = (cell.get("inferText") or cell.get("text") or "")
        s = re.sub(r"\s+", " ", str(s)).strip()
        if s:
            lines.append(s)

    return "\n".join(lines).strip()

# Clova OCR 결과 중에서 페이지의 일반 텍스트 영역(fields)을 전부 이어 붙여 페이지 전체 텍스트로 만듬
def _page_text_from_fields(result_json: dict) -> str:
    img0 = (result_json.get("images") or [{}])[0]
    fields = img0.get("fields") or []
    return " ".join([f.get("inferText", "") for f in fields if f.get("inferText")]).strip()

# 테이블에서 rowSpan / colSpan(병합셀) 정보를 보고, 병합된 셀의 텍스트를 병합 영역 전체 좌표에 복제해서 셀 리스트 만듬
def _merge_cells_with_rowspan(cells: list[dict]) -> list[dict]:
    has_span = any(("rowSpan" in c) or ("colSpan" in c) for c in cells)
    if not has_span:
        return cells

    expanded_cells: list[dict] = []
    for cell in cells:
        row_idx = cell.get("rowIndex", 0)
        col_idx = cell.get("columnIndex", 0)
        row_span = cell.get("rowSpan", 1) or 1
        col_span = cell.get("colSpan", 1) or 1
        text = _cell_text_keep_lines(cell)

        for r in range(row_idx, row_idx + row_span):
            for c in range(col_idx, col_idx + col_span):
                expanded_cells.append(
                    {
                        "rowIndex": r,
                        "columnIndex": c,
                        "text": text,
                        "isMerged": (r != row_idx or c != col_idx),
                    }
                )

    return expanded_cells

# Clova OCR 결과에서 tables를 꺼내서 2차원 배열(grid) 형태로 정리
def _tables_to_schema(result_json: dict) -> dict:
    img0 = (result_json.get("images") or [{}])[0]
    tables = img0.get("tables") or []

    out: dict = {}
    for ti, t in enumerate(tables, start=1):
        cells = t.get("cells") or []
        if not cells:
            out[f"table{ti}"] = {"rows": [], "rowCount": 0, "colCount": 0}
            continue

        has_rc = any("rowIndex" in c for c in cells) and any("columnIndex" in c for c in cells)
        if not has_rc:
            texts = [_cell_text_keep_lines(c) for c in cells]
            texts = [x for x in texts if x]
            rows = [[x] for x in texts]
            out[f"table{ti}"] = {"rows": rows, "rowCount": len(rows), "colCount": 1 if rows else 0}
            continue

        expanded_cells = _merge_cells_with_rowspan(cells)

        max_r = max(c.get("rowIndex", 0) for c in expanded_cells)
        max_c = max(c.get("columnIndex", 0) for c in expanded_cells)
        grid = [[""] * (max_c + 1) for _ in range(max_r + 1)]

        for c in expanded_cells:
            r = c.get("rowIndex", 0)
            col = c.get("columnIndex", 0)
            text = c.get("text", "")
            if not grid[r][col]:
                grid[r][col] = text

        out[f"table{ti}"] = {
            "rows": grid,
            "rowCount": len(grid),
            "colCount": len(grid[0]) if grid else 0,
        }

    return out

# 전체 파이프라인 엔트리 함수
def clovaOCR(file_path: str, lang: str = "ko", dpi: int = 200, enable_table: bool = True) -> dict:
    p = Path(file_path)
    if not p.exists():
        raise FileNotFoundError(str(p.resolve()))

    pages_out: dict = {}

    if p.suffix.lower() == ".pdf":
        page_pngs = _pdf_to_png_bytes_list(str(p), dpi=dpi)

        for i, (fname, png_bytes) in enumerate(page_pngs, start=1):
            print(f"[{i}/{len(page_pngs)}] OCR...") # OCR 되는지 확인 용 디버깅
            rj = _call_clova_from_bytes(
                fname, png_bytes, fmt="png", lang=lang, enable_table=enable_table
            )
            tables = _tables_to_schema(rj)

            pages_out[f"page{i}"] = {
                "text": _page_text_from_fields(rj),
                "tableCount": len(tables),
                "tables": tables,
            }

        return {"pageCount": len(page_pngs), "pages": pages_out}

    with open(p, "rb") as f:
        b = f.read()

    fmt = p.suffix.lower().lstrip(".")
    if fmt == "jpeg":
        fmt = "jpg"

    rj = _call_clova_from_bytes(p.name, b, fmt=fmt, lang=lang, enable_table=enable_table)
    tables = _tables_to_schema(rj)

    pages_out["page1"] = {
        "text": _page_text_from_fields(rj),
        "tableCount": len(tables),
        "tables": tables,
    }
    return {"pageCount": 1, "pages": pages_out}

# 표처리 헬퍼
def clean_table(tbl):   
    return [
        [c for c in ("" if c is None else str(c).strip() for c in row) if c]
        for row in tbl
        if any("" if c is None else str(c).strip() for c in row)
    ]


CSV_EXT = {"csv"}
EXCEL_EXT = {"xls", "xlsx"}
IMAGE_EXT = {"jpg", "png"}
OTHER_EXT = {"pdf", "docx"}

def handle_csv(path: str, slotName: str):
    for enc in ["euc-kr", "cp949", "utf-8"]:
        try:
            df = pd.read_csv(path, encoding=enc)
            return {"slotName": slotName, "kind": "CSV", "ext": "csv", "dataframe": df}
        except UnicodeDecodeError:
            continue
    raise ValueError("인코딩 실패: cp949, euc-kr, utf-8 encoding을 지원합니다")

def handle_excel(path: str, ext: str, slotName: str):
    df = pd.read_excel(path)
    return {"slotName": slotName, "kind": "EXCEL", "ext": ext, "dataframe": df}

def handle_image(path: str, ext: str, slotName: str):
    ocr_output = clovaOCR(path)
    return {"slotName": slotName, "kind": "IMAGE", "ext": ext, "content": ocr_output}

def clean_table(tbl):
    return [
        [c for c in ("" if c is None else str(c).strip() for c in row) if c]
        for row in tbl
        if any("" if c is None else str(c).strip() for c in row)
    ]

def handle_pdf(path: str, ext: str, slotName: str):
    PDF_PAGE_MIN_CHARS = 30
    PDF_PASS_RATIO = 0.7

    with pdfplumber.open(path) as pdf:
        total_pages = len(pdf.pages)
        valid_page_count = 0
        pages_data = {}

        for i, page in enumerate(pdf.pages):
            t = page.extract_text() or ""
            clean_t = " ".join(t.split()).strip()
            if len(clean_t) >= PDF_PAGE_MIN_CHARS: valid_page_count += 1
            raw = page.extract_tables() or []
            tables = {f"table{ti+1}": (lambda rows: {"rows": rows, "rowCount": len(rows), "colCount": max((len(r) for r in rows), default=0)})(clean_table(tbl)) for ti, tbl in enumerate(raw)}
            pages_data[f"page{i+1}"] = {"text": clean_t, "tableCount": len(raw), "tables": tables}

        pass_ratio = valid_page_count / max(total_pages, 1)

    if pass_ratio >= PDF_PASS_RATIO: return {"slotName": slotName, "kind":"PDF","ext":ext,"mode":"text","pass_ratio":pass_ratio,"content":{"pageCount":total_pages,"pages":pages_data}}
    else:
        ocr_output = clovaOCR(path)
        return {"slotName": slotName, "kind": "PDF", "ext": ext, "mode": "ocr", "content": ocr_output}

def handle_docx(path: str, ext: str, slotName: str):
    doc = Document(path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    tables = {}
    for i, table in enumerate(doc.tables, start=1):
        rows = [[cell.text for cell in row.cells] for row in table.rows]
        tables[f"table{i}"] = {
            "rows": rows,
            "rowCount": len(rows),
            "colCount": max((len(r) for r in rows), default=0),
        }
    pages = {
        "page1": {
            "text": "\n".join(paragraphs),
            "tableCount": len(doc.tables),
            "tables": tables,
        }
    }
    return {"slotName": slotName, "kind": "DOCX", "ext": ext, "content": {"pageCount": 1, "pages": pages}}

# 파일 형식별 조건부 함수구동
def handle_file(path: str, slotName: str):
    ext = Path(path).suffix.lower().lstrip(".")
    if ext in CSV_EXT: return handle_csv(path, slotName)
    if ext in EXCEL_EXT: return handle_excel(path, ext, slotName)
    if ext in IMAGE_EXT: return handle_image(path, ext, slotName)
    if ext == "pdf": return handle_pdf(path, ext, slotName)
    if ext == "docx": return handle_docx(path, ext, slotName)
    raise ValueError(f"지원하지 않는 파일 형식입니다: {ext}")